"""
Document Parser - Extracts text from various file formats.

Supported formats:
- PDF (.pdf)
- Word (.docx)
- Plain text (.txt)

Design decisions:
- Preserve document structure (sections, paragraphs)
- Extract metadata (title, author, dates)
- Handle encoding issues gracefully
- Clean up extracted text
"""

from pathlib import Path
from typing import Optional, Dict, Any
from enum import Enum
import re

from pydantic import BaseModel, Field
from loguru import logger

# Document parsing libraries
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf not installed - PDF parsing unavailable")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed - DOCX parsing unavailable")


class DocumentType(str, Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    UNKNOWN = "unknown"


class ParsedDocument(BaseModel):
    """
    Parsed document with extracted text and metadata.
    
    Why structured this way:
    - text: Full document text (for chunking)
    - metadata: Document info (for filtering, audit trails)
    - document_type: For type-specific processing
    """
    text: str = Field(..., description="Extracted document text")
    metadata: Dict[str, Any] = Field(
        default_factory=dict,
        description="Document metadata (title, author, page count, etc.)"
    )
    document_type: DocumentType = Field(..., description="Type of document")
    file_path: str = Field(..., description="Original file path")
    
    class Config:
        json_schema_extra = {
            "example": {
                "text": "EMPLOYMENT AGREEMENT\n\nSection 12.3 - Dispute Resolution...",
                "metadata": {
                    "title": "Employment Agreement",
                    "pages": 15,
                    "author": "Legal Department"
                },
                "document_type": "pdf",
                "file_path": "/path/to/document.pdf"
            }
        }


class DocumentParser:
    """
    Parses documents from various formats into structured text.
    
    Usage:
        parser = DocumentParser()
        doc = parser.parse_file("contract.pdf")
        print(doc.text)
    
    Why separate parser:
    - Centralized document handling
    - Easy to add new file types
    - Consistent error handling
    - Testable in isolation
    """
    
    @staticmethod
    def detect_document_type(file_path: Path) -> DocumentType:
        """
        Detect document type from file extension.
        
        Why explicit detection:
        - Don't rely on MIME types (can be spoofed)
        - File extension is good enough for our use case
        - Easy to add new types
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return DocumentType.PDF
        elif suffix in [".docx", ".doc"]:
            return DocumentType.DOCX
        elif suffix == ".txt":
            return DocumentType.TXT
        else:
            return DocumentType.UNKNOWN
    
    def parse_file(self, file_path: str | Path) -> ParsedDocument:
        """
        Parse a document file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            ParsedDocument with extracted text and metadata
            
        Raises:
            ValueError: If file type not supported
            FileNotFoundError: If file doesn't exist
        
        Learning insight:
        - Different file types need different parsing strategies
        - Always validate file exists before parsing
        - Extract metadata for better retrieval/filtering
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        doc_type = self.detect_document_type(file_path)
        
        logger.info(f"Parsing {doc_type.value} document: {file_path.name}")
        
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(file_path)
        elif doc_type == DocumentType.TXT:
            return self._parse_txt(file_path)
        else:
            raise ValueError(
                f"Unsupported document type: {file_path.suffix}. "
                f"Supported: .pdf, .docx, .txt"
            )
    
    def _parse_pdf(self, file_path: Path) -> ParsedDocument:
        """
        Parse PDF document using pypdf.
        
        Why pypdf:
        - Lightweight (no dependencies on external tools)
        - Good text extraction for most PDFs
        - Handles metadata extraction
        
        Limitations:
        - OCR not included (scanned PDFs won't work)
        - Complex layouts may have extraction issues
        - For production, consider pdfplumber or OCR pipeline
        """
        if not PYPDF_AVAILABLE:
            raise ImportError(
                "pypdf not installed. Install with: pip install pypdf"
            )
        
        try:
            reader = pypdf.PdfReader(file_path)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Clean up text (remove extra whitespace, normalize)
            full_text = self._clean_text(full_text)
            
            # Extract metadata
            metadata = {
                "pages": len(reader.pages),
                "file_size_bytes": file_path.stat().st_size,
            }
            
            # Try to extract PDF metadata (may not always be present)
            if reader.metadata:
                if reader.metadata.title:
                    metadata["title"] = reader.metadata.title
                if reader.metadata.author:
                    metadata["author"] = reader.metadata.author
                if reader.metadata.subject:
                    metadata["subject"] = reader.metadata.subject
                if reader.metadata.creator:
                    metadata["creator"] = reader.metadata.creator
            
            logger.success(
                f"Extracted {len(full_text)} chars from {len(reader.pages)} pages"
            )
            
            return ParsedDocument(
                text=full_text,
                metadata=metadata,
                document_type=DocumentType.PDF,
                file_path=str(file_path)
            )
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
    
    def _parse_docx(self, file_path: Path) -> ParsedDocument:
        """
        Parse DOCX document using python-docx.
        
        Why python-docx:
        - Native DOCX support (no conversion needed)
        - Preserves document structure
        - Handles tables, formatting
        
        Limitations:
        - Only .docx (not old .doc format)
        - Complex formatting may be lost
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx not installed. Install with: pip install python-docx"
            )
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Clean up text
            full_text = self._clean_text(full_text)
            
            # Extract metadata
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "file_size_bytes": file_path.stat().st_size,
            }
            
            # Try to extract core properties
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject
            
            logger.success(
                f"Extracted {len(full_text)} chars from {len(doc.paragraphs)} paragraphs"
            )
            
            return ParsedDocument(
                text=full_text,
                metadata=metadata,
                document_type=DocumentType.DOCX,
                file_path=str(file_path)
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
    
    def _parse_txt(self, file_path: Path) -> ParsedDocument:
        """
        Parse plain text document.
        
        Why support TXT:
        - Simplest format (no parsing needed)
        - Good for testing
        - Some contracts are plain text
        
        Encoding handling:
        - Try UTF-8 first (most common)
        - Fall back to latin-1 if UTF-8 fails
        - log warning if had to fall back
        """
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                # Fall back to latin-1
                logger.warning(f"UTF-8 decode failed for {file_path}, trying latin-1")
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
            
            # Clean up text
            text = self._clean_text(text)
            
            # Basic metadata
            metadata = {
                "file_size_bytes": file_path.stat().st_size,
                "lines": text.count('\n') + 1,
            }
            
            logger.success(f"Extracted {len(text)} chars from TXT file")
            
            return ParsedDocument(
                text=text,
                metadata=metadata,
                document_type=DocumentType.TXT,
                file_path=str(file_path)
            )
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            raise
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean up extracted text.
        
        Cleaning steps:
        1. Normalize whitespace (collapse multiple spaces)
        2. Fix line breaks (remove orphaned breaks)
        3. Remove control characters
        4. Trim leading/trailing whitespace
        
        Why clean:
        - PDFs often have weird formatting artifacts
        - Improves chunking quality
        - Reduces noise in embeddings
        
        Design decision: Conservative cleaning
        - Don't remove too much (preserve structure)
        - Legal text needs precision
        """
        # Remove null bytes and control characters (except newlines/tabs)
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize line breaks (but preserve paragraph breaks)
        text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
        
        # Collapse multiple spaces (but preserve indentation)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Final trim
        text = text.strip()
        
        return text
